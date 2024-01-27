##PARA CORRER O PROGRAMA python3 main.py

from time import sleep
from datetime import datetime
from time import sleep
from docx import Document
from flask import Blueprint, render_template, request, redirect, url_for, jsonify, send_file
from . import views
from .data_base import *
import os, shutil, zipfile, re

views = Blueprint("views", __name__)

#TURMAS
@views.route("/turmas", methods=["GET", "POST"])
def turmas():
    turmasData = session.query(Turmas).all()
    return render_template("turmas.html", turmasData=turmasData)

@views.route("/criar_turma", methods=["GET", "POST"])
def criar_turma():
    if request.method == "POST":
        turma = request.form["turma"]
        turmaInserir = Turmas(descricao=turma)
        turma = session.query(Turmas).filter_by(descricao=turma).first()
        arr = []
        if turma == None:
            arr.append(0)
            session.add(turmaInserir)
            session.commit()
        else:
            arr.append("Turma já criada")
        return jsonify(arr)
    return render_template("criar_turma.html")

@views.route("/editar_turma/<int:id>", methods=["GET", "POST"])
def editar_turma(id):
    if request.method == "POST":
        turma_desc = request.form["turma"]
        turma = session.query(Turmas).filter_by(descricao=turma_desc).first()
        turma_antiga = session.query(Turmas).filter_by(id=id).first()
        arr = []
        if turma != None:
            if turma.descricao == turma_antiga.descricao:
                arr.append("Não pode alterar o nome da turma para o mesmo nome")
                return jsonify(arr)
        if turma == None:
            arr.append(0)
            turma_antiga.descricao = turma_desc
            session.commit()
        else:
            arr.append("Já existe uma turma com o mesmo nome")
        return jsonify(arr)
    turmaEditar = session.query(Turmas).get(id)
    if turmaEditar == None:
        return redirect(url_for("views.turmas"))
    return render_template("editar_turma.html", turmaEditar=turmaEditar, id=id)

@views.route("/eliminar_turma/<int:id>", methods=["GET", "POST"])
def eliminar_turma(id):
    turma = session.query(Turmas).filter_by(id=id).first()
    if turma == None:
        return redirect(url_for("views.turmas"))
    session.delete(turma)
    session.commit()
    return redirect(url_for("views.turmas"))

#ALUNOS
@views.route("/alunos", methods=["GET", "POST"])
def alunos():
    turmasData = session.query(Turmas).all()
    return render_template("alunos.html", turmasData=turmasData)

@views.route("/criar_aluno", methods=["GET", "POST"])
def criar_aluno():
    turmasData = session.query(Turmas).all()
    if request.method == "POST":
        nif = request.form["nif"]
        nr = request.form["numero"]
        turmaDesc = request.form["inputTurma"]
        turma = session.query(Turmas).filter_by(descricao=turmaDesc).first()
        turma_id = turma.id
        nome = request.form["nome"]
        nome_split = nome.split()
        nome_abreviado = nome_split[0] + " " + nome_split[-1]
        adress = request.form["morada"]
        codPostal = request.form["codpostal"]
        localidade = request.form["local"]
        codPostal = codPostal + " " + localidade
        validade = request.form["validade"]
        cc = request.form["cc"]
        alunoInserir = Alunos(nif=nif,nr=nr,turmaId=turma_id,nome=nome,morada=adress,cod_postal=codPostal,validade_cc=validade,cartao_cidadao=cc,nome_abreviado=nome_abreviado,)
        arr = []
        if session.query(Alunos).filter_by(turmaId=turma_id, nr=nr).first() is None:
            if session.query(Alunos).filter_by( nif=nif).first() is None:
                if session.query(Alunos).filter_by(cartao_cidadao=cc).first() is None:
                    pass
                else:    
                    arr.append("Já existe um aluno com o mesmo CC")
            else:
                arr.append("Já existe um aluno com o mesmo NIF")
        else:
            arr.append( "Aluno com o mesmo número já criado")
        
        if len(arr) != 0:
            pass
        else:
            arr.append(0)
            session.add(alunoInserir)        
            session.commit()
        return jsonify(arr), 200
    return render_template("criar_aluno.html", turmasData = turmasData)

@views.route("/editar_aluno/<int:id>", methods=["GET", "POST"])
def editar_aluno(id):
    if request.method == "POST":
        aluno = session.query(Alunos).filter_by(id=id).first()
        nif = request.form["nif"]
        nr = request.form["numero"]
        turmaId = request.form["inputTurma"]
        turma = session.query(Turmas).filter_by(id=turmaId).first()
        turma_id = turma.id
        nome = request.form["nome"]
        nome_split = nome.split()
        nome_abreviado = nome_split[0] + " " + nome_split[-1]
        adress = request.form["morada"]
        codPostal = request.form["codpostal"]
        localidade = request.form["local"]
        codPostal = codPostal + " " + localidade
        validade = request.form["validade"]
        cc = request.form["cc"]
        arr = []
        nraluno = session.query(Alunos).filter_by(turmaId=turma_id, nr=nr).first()
        nifaluno = session.query(Alunos).filter_by( nif=nif).first()
        ccaluno = session.query(Alunos).filter_by(cartao_cidadao=cc).first()
        if nraluno is None:
            pass
        else:
            if(nraluno.id != id):
                arr.append( "Aluno com o mesmo número já criado")
        if nifaluno is None:
            pass
        else:
            if(nifaluno.id != id):
                arr.append("Já existe um aluno com o mesmo NIF")
        if ccaluno is None:
            pass
        else:
            if(ccaluno.id != id):
                arr.append( "Aluno com o mesmo cc já criado")

        if len(arr) != 0:
            pass
        else:
            arr.append(0)
            aluno.nif = nif
            aluno.nr = nr
            aluno.turmaId = turma_id
            aluno.nome = nome 
            aluno.nome_abreviado = nome_abreviado
            aluno.morada = adress
            aluno.cod_postal = codPostal
            aluno.localidade = localidade
            aluno.validade_cc = validade
            aluno.cartao_cidao = cc
            session.commit()
        return jsonify(arr), 200
    turmasData = session.query(Turmas).all()
    alunoEditar = session.query(Alunos).get(id)
    if alunoEditar == None:
        return redirect(url_for("views.alunos"))
    localidade = alunoEditar.cod_postal[8:].strip()
    print(localidade)
    alunoEditar.cod_postal = alunoEditar.cod_postal[:8]
    return render_template("editar_aluno.html", alunoEditar=alunoEditar, id=id, turmasData=turmasData, localidade = localidade)

@views.route("/eliminar_aluno/<int:id>", methods=["GET", "POST"])
def eliminar_aluno(id):
    aluno = session.query(Alunos).filter_by(id=id).first()
    if aluno is None:
        return redirect(url_for("views.alunos"))
    session.delete(aluno)
    session.commit()
    return redirect(url_for("views.alunos"))






#ENTIDADES
@views.route("/entidades")
def entidades():
    entidadesData = session.query(Entidade).all()
    return render_template("entidades.html", entidadesData=entidadesData)
@views.route("/criar_entidade", methods=["GET", "POST"])
def criar_entidade():
    if request.method == "POST":
        nif = request.form["nif"]
        nome = request.form["nome"]
        adress = request.form["morada"]
        codPostal = request.form["cod_postal"]
        pessoa_responsavel = request.form["pessoa_responsavel"]
        cargo_responsavel = request.form["cargo_responsavel"]
        localidade = request.form["localidade"]
        codPostal = codPostal + " " + localidade
        entidades = Entidade(nif=nif,nome=nome,morada=adress,cod_postal=codPostal,cargo_pessoa_responsavel=cargo_responsavel,pessoa_responsavel=pessoa_responsavel,)
        arr = []
        if session.query(Entidade).filter_by(nif = nif).first() is None:
            if session.query(Entidade).filter_by(nome = nome).first() is None:
                pass
            else:
                arr.append("Entidade com o mesmo nome já adicionada")
        else:
                arr.append("Entidade com o mesmo NIF já adicionada")
        if len(arr) != 0:
            pass
        else:
            arr.append(0)
            session.add(entidades)
            session.commit()    
        return jsonify(arr), 200
    return render_template("criar_entidade.html")

@views.route("/editar_entidade/<int:id>", methods=["GET", "POST"])
def editar_entidade(id):
    if request.method == "POST":
        entidade = session.query(Entidade).filter_by(id=id).first()
        nif = request.form["nif"]
        nome = request.form["nome"]
        morada = request.form["morada"]
        cod_postal = request.form["cod_postal"]
        pessoa_responsavel = request.form["pessoa_responsavel"]
        cargo_pessoa_responsavel = request.form["cargo_responsavel"]
        localidade = request.form["localidade"]
        cod_postal = cod_postal + " " + localidade
        arr = []
        nifEntidade = session.query(Entidade).filter_by(nif = nif).first()
        nomeEntidade = session.query(Entidade).filter_by(nome = nome).first()
        if nifEntidade is None:
            pass
        else:
            if nifEntidade.id != id:
                arr.append( "Entidade com o mesmo nif já adicionada")
        if nomeEntidade is None:
            pass
        else: 
            if nomeEntidade.id != id:
                arr.append("Entidade já com o mesmo nome criada")
                
        if len(arr) != 0:
            pass
        else:
            arr.append(0)
            entidade.nif = nif
            entidade.nome = nome
            entidade.morada = morada
            entidade.cod_postal = cod_postal
            entidade.pessoa_responsavel = pessoa_responsavel
            entidade.cargo_pessoa_responsavel = cargo_pessoa_responsavel
            session.commit()   
        return jsonify(arr), 200
    entidadeEditar = session.query(Entidade).get(id)
    if entidadeEditar == None:
        return redirect(url_for("views.entidades"))
    localidade = entidadeEditar.cod_postal[8:].strip()
    entidadeEditar.cod_postal = entidadeEditar.cod_postal[:8]
    return render_template("editar_entidade.html", entidadeEditar=entidadeEditar, id=id, localidade = localidade)
    
@views.route("/eliminar_enditade/<int:id>", methods=["GET", "POST"])
def eliminar_entidade(id):
    entidade = session.query(Entidade).filter_by(id=id).first()
    if entidade is None:
        return redirect(url_for("views.entidades"))
    session.delete(entidade)
    session.commit()
    return redirect(url_for("views.entidades"))



#ESTAGIOS
@views.route("/")
@views.route("/estagios")
def estagios():
    alunosData = []
    turmasData = session.query(Turmas).all()
    estagiosData = session.query(Estagios).all()
    for i in estagiosData:
        aluno = session.query(Alunos).filter_by(id=i.alunoId).first()
        alunosData.append(aluno)
    data = zip(estagiosData, alunosData)     
    return render_template("estagios.html", data=data, estagiosData = estagiosData, turmasData = turmasData)

@views.route("/criar_estagio", methods=["GET", "POST"])
def criar_estagio():
    turmasData = session.query(Turmas).all()
    entidadesData = session.query(Entidade).all()
    if request.method == "POST":
        meses = {'Janeiro': 1,'Fevereiro': 2,'Março': 3,'Abril': 4,'Maio': 5,'Junho': 6,'Julho': 7,'Agosto': 8,'Setembro': 9,'Outubro': 10,'Novembro': 11,'Dezembro': 12}
        alunoId = request.form["aluno"]
        entidade = request.form["entidade"]
        diaInicio = request.form["inputDia"]
        diaFim = request.form.get("inputDiaFim")
        mesInicio = request.form["inputMes"]
        mesFim = request.form.get("inputMesFim")
        anoIncio = request.form["inputAno"]
        anoFim = request.form.get("inputAnoFim")
        morada = request.form["morada"]
        cod_postal = request.form["cod_postal"]
        tutor = request.form["tutor"]
        orientador = request.form["orientador"]
        email_tutor = request.form["email_tutor"]
        entidadeId = session.query(Entidade).filter_by(id=entidade).first().id
        data_inicio = str(anoIncio) + "-" + str(meses[mesInicio]) + "-" + str(diaInicio) 
        data_fim = str(anoFim) + "-" + str(meses[mesFim]) + "-" + str(diaFim)
        isAlunoId = session.query(Estagios).filter_by(alunoId = alunoId).first()
        localidade = request.form["localidade"]
        cod_postal = cod_postal + " " + localidade
        arr = []
        email_is_valido = verificar_email(email_tutor)
        if email_is_valido == False:
            arr.append("Formato de email inválido")
            return jsonify(arr)    

        if isAlunoId == None:
            arr.append(0)
            estagio = Estagios(alunoId=alunoId,morada=morada, cod_postal=cod_postal,data_inicio = data_inicio,data_fim = data_fim,tutor=tutor,email_tutor=email_tutor,entidadeId=entidadeId,orientador=orientador,)
            session.add(estagio)
            session.commit()
        else:
            arr.append("Aluno já com um estágio")
        return jsonify(arr)    
    return render_template("criar_estagio.html",turmasData=turmasData,entidadesData=entidadesData)

def verificar_email(email):
    padrao = r'^[\w\.-]+@[\w\.-]+\.\w+$'
    if re.match(padrao, email):
        return True
    else:
        return False

@views.route("/editar_estagio/<int:id>", methods=["GET", "POST"])
def editar_estagio(id):
    estagioEditar = session.query(Estagios).filter_by(id = id).first()
    if estagioEditar == None:
        return redirect(url_for("views.estagios"))
    turmasData = session.query(Turmas).all()
    entidadesData = session.query(Entidade).all()
    if request.method == "POST":
        meses = {'Janeiro': 1,'Fevereiro': 2,'Março': 3,'Abril': 4,'Maio': 5,'Junho': 6,'Julho': 7,'Agosto': 8,'Setembro': 9,'Outubro': 10,'Novembro': 11,'Dezembro': 12}
        alunoId = request.form["aluno"]
        entidade = request.form["entidade"]
        print(entidade)
        diaInicio = request.form["inputDia"]
        diaFim = request.form.get("inputDiaFim")
        mesInicio = request.form["inputMes"]
        mesFim = request.form.get("inputMesFim")
        anoIncio = request.form["inputAno"]
        anoFim = request.form.get("inputAnoFim")
        morada = request.form["morada"]
        cod_postal = request.form["cod_postal"]
        tutor = request.form["tutor"]
        orientador = request.form["orientador"]
        email_tutor = request.form["email_tutor"]
        localidade = request.form["localidade"]
        cod_postal = cod_postal + " " + localidade
        data_inicio = str(anoIncio) + "-" + str(meses[mesInicio]) + "-" + str(diaInicio) 
        data_fim = str(anoFim) + "-" + str(meses[mesFim]) + "-" + str(diaFim)
        isAlunoId = session.query(Estagios).filter_by(alunoId = alunoId).first()
        arr = []
        email_is_valido = verificar_email(email_tutor)
        if email_is_valido == False:
            arr.append("Formato de email inválido")
            return jsonify(arr)
        if isAlunoId == None:
            arr.append(0)
            estagio = session.query(Estagios).filter_by(id = id).first()
            estagio.alunoId = alunoId
            estagio.entidadeId = entidade
            estagio.data_inicio = data_inicio
            estagio.data_fim = data_fim
            estagio.morada = morada
            estagio.cod_postal = cod_postal
            estagio.tutor = tutor
            estagio.orientador = orientador
            estagio.email_tutor = email_tutor
            session.commit()
        else:
            if isAlunoId.id == id:
                arr.append(0)
                estagio = session.query(Estagios).filter_by(id = id).first()
                estagio.alunoId = alunoId
                estagio.entidadeId = entidade
                estagio.data_inicio = data_inicio
                estagio.data_fim = data_fim
                estagio.morada = morada
                estagio.cod_postal = cod_postal
                estagio.tutor = tutor
                estagio.orientador = orientador
                estagio.email_tutor = email_tutor
                session.commit()
            else:
                arr.append("Aluno já com um estágio ativo, elemie o estágio desse aluno ou altere o mesmo")
        return jsonify(arr)
    localidade = estagioEditar.cod_postal[8:]
    codPostal = estagioEditar.cod_postal[:8]
    aluno = session.query(Alunos).filter_by(id = estagioEditar.alunoId).first()
    nomeTurma = session.query(Turmas).filter_by(id = aluno.turmaId).first().descricao

    return render_template("editar_estagio.html", estagioEditar=estagioEditar, id=id, turmasData=turmasData,entidadesData=entidadesData, nomeTurma = nomeTurma, localidade = localidade, codPostal = codPostal)


@views.route("/eliminar_estagio/<int:id>", methods=["GET", "POST"])
def eliminar_estagio(id):
    estagio = session.query(Estagios).filter_by(id=id).first()
    if estagio is None:
        return redirect(url_for("views.estagios"))
    session.delete(estagio)
    session.commit()
    return redirect(url_for("views.estagios"))
        
def gerar_doc(document, placeholders_dict):
    for paragraph in document.paragraphs:
        for key, value in placeholders_dict.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in placeholders_dict.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

def zip_directory(directory_path, zip_path):
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, _, files in os.walk(directory_path):
            for file in files:
                file_path = os.path.join(root, file)
                zipf.write(file_path, os.path.relpath(file_path, directory_path))

@views.route('/download/<int:id>')
def documentos(id):
    estagio = session.query(Estagios).filter_by(id=id).first()
    if estagio is None:
        return redirect(url_for("views.estagios"))
    aluno = session.query(Alunos).filter_by(id=estagio.alunoId).first()
    turma = session.query(Turmas).filter_by(id=aluno.turmaId).first()
    entidade = session.query(Entidade).filter_by(id=estagio.entidadeId).first()
    nome = aluno.nome_abreviado.strip()
    nome = nome.replace(" ", "_")
    downloads_path = f'Downloads/{nome}'
    data = str(datetime.now().strftime("%d/%m/%Y"))
    placeholders_dict = {"[nome]": aluno.nome,"[turma]": turma.descricao,"[alunoCartaoC]": aluno.cartao_cidadao,"[AVcc]": aluno.validade_cc.strftime("%d/%m/%Y"),"[Anif]": aluno.nif,"[Amorada]": aluno.morada,"[nr]": aluno.nr,"[tutor]": estagio.tutor,"[orientador]": estagio.orientador,"[Enome]": entidade.nome,"[ENIF]": entidade.nif,"[Emorada]": entidade.morada,"[EPrespons]": entidade.pessoa_responsavel,"[ECPresponsavel]": entidade.cargo_pessoa_responsavel,"[data_inicio]": estagio.data_inicio.strftime("%d/%m/%Y"),"[data_fim]": estagio.data_fim.strftime("%d/%m/%Y"),"[data]" : data}

    document_path_1 = 'docs/folha_de_presencas.docx'
    document_path_2 = 'docs/plano_de_fct.docx'
    document_path_3 = 'docs/plano_individual_de_fct.docx'
    document_path_4 = 'docs/grelha_de_avaliacao.docx'
    document_path_5 = 'docs/Protocolo_de_FCT.docx'
    
    if not os.path.exists(downloads_path):
        os.makedirs(downloads_path)
    
    document1 = Document(document_path_1)
    gerar_doc(document1, placeholders_dict)
    new_document_path_1 = os.path.join(downloads_path, 'Imp. PEA-3.4-002-0 Folha de Presenças e Registo de Atividades Semanal.docx.docx')
    document1.save(new_document_path_1)

    document2 = Document(document_path_2)
    gerar_doc(document2, placeholders_dict)
    new_document_path_2 = os.path.join(downloads_path, 'Imp. PEA-3.4-003-0 Plano de FCT.docx')
    document2.save(new_document_path_2)
    
    document3 = Document(document_path_3)
    gerar_doc(document3, placeholders_dict)
    new_document_path_3 = os.path.join(downloads_path, 'Imp. PEA-3.4-011-0 Plano Individual de FCT.docx')
    document3.save(new_document_path_3)
    
    document4 = Document(document_path_4)
    gerar_doc(document4, placeholders_dict)
    new_document_path_4 = os.path.join(downloads_path, 'Imp. PEA-3.4-005-0 Grelha de Avaliação da FCT.docx')
    document4.save(new_document_path_4)
    
    document5 = Document(document_path_5)
    gerar_doc(document5, placeholders_dict)
    new_document_path_5 = os.path.join(downloads_path, 'Imp. PEA-3.4-010-2 Protocolo de FCT.docx')    
    document5.save(new_document_path_5)


    nome_arquivo_zip = f"{nome}.zip"
    nome_pasta = f"{nome}"
    caminho_diretorio = os.path.join(os.getcwd(), "Downloads")
    zip_directory(caminho_diretorio+f"/{nome_pasta}", caminho_diretorio+"/"+nome_arquivo_zip)
    zip_path = caminho_diretorio+"/"+nome_arquivo_zip
    return send_file(zip_path, as_attachment=True)

@views.after_request
def remove_file(response):
    download_folder = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'Downloads')
    if os.path.exists(download_folder) and os.path.isdir(download_folder):
        for item in os.listdir(download_folder):
            file_path = os.path.join(download_folder, item)
            try:
                if os.path.isfile(file_path):
                    # Remove o arquivo
                    os.remove(file_path)
                elif os.path.isdir(file_path):
                    # Remove o diretório (e seu conteúdo) recursivamente
                    shutil.rmtree(file_path)
            except Exception as e:
                    print(f"Erro ao excluir {file_path}: {e}")
    return response

@views.route("/get_dia", methods=["GET", "POST"])
def get_dia():
        dias = []
        dia_atual = datetime.today().day        
        data = request.get_json()
        ano = data['anoSelecionado']
        mes = data['mesSelecionado'] 
        if ano == 1 and mes == 1:
            pass
        if ano == 0 and mes == 0:
            mes_atual = datetime.today().month
            ano_atual = datetime.today().year
            anos = list(range(1900, ano_atual+5))
            mes_atual = {1: "Janeiro", 2: "Fevereiro", 3: "Março", 4: "Abril", 5: "Maio", 6: "Junho", 7: "Julho", 8: "Agosto", 9: "Setembro", 10: "Outubro", 11: "Novembro", 12: "Dezembro"}.get(mes_atual)
            meses = ["Janeiro", "Fevereiro", "Março", "Abril", "Maio", "Junho","Julho", "Agosto", "Setembro", "Outubro", "Novembro", "Dezembro"]
            dict = {"type": 0, "Meses": meses, "Anos": anos, "Dia": dia_atual, "Mes": mes_atual, "Ano": ano_atual}
            return jsonify(dict)
        if mes == "Janeiro" or mes == "Março" or mes == "Maio" or mes == "Julho" or mes == "Agosto" or mes == "Outubro" or mes == "Dezembro":
            dias = list(range(1, 32))
        elif mes == "Abril" or mes == "Junho" or mes == "Setembro" or mes == "Novembro":
            dias = list(range(1, 31))
        elif mes == "Fevereiro":
            ano = int(ano)
            if (ano%4==0 and ano%100!=0) or (ano%400==0):
                dias = list(range(1, 30))
            else:
                dias = list(range(1, 29))
        return jsonify(dias, dia_atual)

@views.route("/get_alunos", methods=["GET", "POST"])
def get_aluno():
    if request.method == "POST":
        aluno_data = []
        data = request.get_json()
        turmaSelecionada = data.get("turmaSelecionada")
        estagioId = data.get("estagioId")
        alunoEstagio = None
        if estagioId is not None:
            estagio = session.query(Estagios).filter_by(id=estagioId).first().alunoId
            alunoEstagio = session.query(Alunos).filter_by(id=estagio).first().nome_abreviado
        alunos = session.query(Alunos).filter_by(turmaId=turmaSelecionada).all()
        for aluno in alunos:
            estagioId = session.query(Estagios).filter_by(alunoId=id).first()
            alunso_data = {'id': aluno.id,'nr': aluno.nr,'nome': aluno.nome,'morada': aluno.morada,'cod_postal': aluno.cod_postal,'cartao_cidadao': aluno.cartao_cidadao,'validade_cc': aluno.validade_cc,'nif': aluno.nif,'nome_abreviado': aluno.nome_abreviado,'estagioId' : estagioId}
            aluno_data.append(alunso_data)
        if alunoEstagio is not None:
            return jsonify(aluno_data, alunoEstagio)
    return jsonify(aluno_data)

@views.route("/get_entidade", methods=["GET", "POST"])
def get_entidade():
    if request.method == "POST":
        entidade_data = {}
        data = request.get_json()
        id = data.get("id")
        id = int(id)
        entidade = session.query(Entidade).filter_by(id=id).first()
        entidade_data["codPostal"] = entidade.cod_postal[:8]
        entidade_data["Morada"] = entidade.morada
        entidade_data["localidade"] = entidade.cod_postal[8:].strip()
        entidade_data["id"] = entidade.id
    return jsonify(entidade_data)

@views.route("/get_alunos_estagios", methods=["GET", "POST"])
def get_aluno_estagio():
    if request.method == "POST":
        estagios_data = []
        data = request.get_json()
        turmaSelecionada = data.get("turmaSelecionada")
        alunos = session.query(Alunos).filter_by(turmaId=turmaSelecionada).all()
        estagios = session.query(Estagios).all()
        alunosId = []
        for alunos in alunos:
            alunosId.append(alunos.id)
        for estagios in estagios:
            if estagios.alunoId in alunosId:
                estagio_data = {'id': estagios.id,'alunoId': estagios.alunoId,'nome': session.query(Alunos).filter_by(id=estagios.alunoId).first().nome,'entidade': session.query(Entidade).filter_by(id=estagios.entidadeId).first().nome,'data_inicio': str(estagios.data_inicio),'data_fim': str(estagios.data_fim),}
                estagios_data.append(estagio_data)
    return jsonify(estagios_data)